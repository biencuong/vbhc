"""Validate a .docx file against the 9 thể thức components of NĐ 30/2020/NĐ-CP.

Usage:
    python validate_thethuc.py <file.docx>

Output: a checklist with ✓ / ⚠ / ✗ per item, plus a summary line.

Note: this is a heuristic check. Manual review is still required for things
like dấu, chữ ký, font sizes (python-docx doesn't reliably read font sizes
from style inheritance).
"""
import re
import sys
from pathlib import Path

from _common import slugify_vn  # noqa: F401
from docx import Document

OK = "✓"
WARN = "⚠"
FAIL = "✗"

PLACEHOLDER_RE = re.compile(r"\?\?\?|<[^>]+>|\[placeholder\]", re.IGNORECASE)


def collect_all_text(doc):
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def check_quoc_hieu(text):
    if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text and "Độc lập" in text and "Tự do" in text and "Hạnh phúc" in text:
        return OK, "Quốc hiệu + Tiêu ngữ"
    return FAIL, "Quốc hiệu + Tiêu ngữ — thiếu hoặc sai"


def check_co_quan(text):
    upper_lines = [ln for ln in text.split("\n") if ln.strip() and ln == ln.upper() and len(ln) > 4]
    if any(re.search(r"(UBND|ỦY BAN|BỘ|SỞ|HĐND|VĂN PHÒNG|TỔNG CỤC|CỤC|CHI CỤC|VIỆN|BAN)", ln) for ln in upper_lines):
        return OK, "Tên cơ quan ban hành"
    return WARN, "Tên cơ quan — không phát hiện được dòng tên cơ quan in hoa"


def is_bieu_mau_noi_bo(text):
    """Detect biểu mẫu nội bộ — Phiếu biểu quyết / Phiếu ghi ý kiến / Phiếu thẩm định.
    Loại này KHÔNG có Số VB, KHÔNG có Nơi nhận, vì là form gửi nội bộ."""
    return bool(re.search(r"PHIẾU\s+(GHI\s+Ý\s+KIẾN|BIỂU\s+QUYẾT|THẨM\s+ĐỊNH|LẤY\s+Ý\s+KIẾN)", text))


def check_so_van_ban(text):
    if is_bieu_mau_noi_bo(text):
        return OK, "Số VB — không cần (biểu mẫu nội bộ)"
    m = re.search(r"Số:\s*([0-9]*\s*/[A-ZĐa-z&\-_]+)", text)
    if m:
        sotxt = m.group(1).strip()
        if sotxt.startswith("/"):
            return WARN, f"Số văn bản — đang trống ({m.group(0).strip()}) — VPHC sẽ điền sau"
        return OK, f"Số văn bản: {m.group(0).strip()}"
    if "Số:" in text:
        return WARN, "Có dòng 'Số:' nhưng không match format chuẩn"
    return FAIL, "Không tìm thấy 'Số:' — kiểm tra lại"


def check_ten_loai(text):
    keywords = ["QUYẾT ĐỊNH", "NGHỊ QUYẾT", "BÁO CÁO", "TỜ TRÌNH", "THÔNG BÁO",
                "KẾ HOẠCH", "CHỈ THỊ", "BIÊN BẢN", "GIẤY MỜI", "PHIẾU GHI Ý KIẾN",
                "PHIẾU BIỂU QUYẾT", "HƯỚNG DẪN", "KẾT LUẬN"]
    found = [k for k in keywords if k in text]
    if found:
        return OK, f"Tên loại: {found[0]}"
    if "V/v" in text or "V/V" in text:
        return OK, "Tên loại: Công văn (có trích yếu V/v...)"
    return WARN, "Không xác định được loại VB"


def check_noi_dung(text):
    if PLACEHOLDER_RE.search(text):
        matches = PLACEHOLDER_RE.findall(text)[:3]
        return FAIL, f"Còn placeholder: {matches}"
    return OK, "Nội dung — không còn placeholder"


def check_nguoi_ky(text):
    upper_chuc_vu = re.search(r"(KT\.|TL\.|TUQ\.)?\s*(GIÁM ĐỐC|CHỦ TỊCH|PHÓ CHỦ TỊCH|CHÁNH VĂN PHÒNG|PHÓ CHÁNH VĂN PHÒNG|TRƯỞNG PHÒNG|PHÓ GIÁM ĐỐC|BỘ TRƯỞNG|THỨ TRƯỞNG|VỤ TRƯỞNG|CỤC TRƯỞNG|VIỆN TRƯỞNG)", text)
    if upper_chuc_vu:
        return OK, f"Chức vụ người ký: {upper_chuc_vu.group(0).strip()}"
    return WARN, "Không phát hiện được chức vụ người ký in hoa"


def check_dau():
    return WARN, "Dấu / chữ ký số — KHÔNG kiểm tra tự động được, kiểm tra thủ công"


def check_noi_nhan(text):
    if is_bieu_mau_noi_bo(text):
        return OK, "Nơi nhận — không cần (biểu mẫu nội bộ)"
    if "Nơi nhận:" in text or "Nơi nhận :" in text:
        if "Lưu:" in text or "Lưu :" in text:
            return OK, "Nơi nhận + Lưu"
        return WARN, "Có 'Nơi nhận' nhưng thiếu 'Lưu:'"
    return FAIL, "Thiếu 'Nơi nhận:'"


def check_phu_luc(text):
    has_kem_theo = bool(re.search(r"kèm theo|đính kèm", text, re.IGNORECASE))
    has_phu_luc = bool(re.search(r"PHỤ LỤC\s+[IVX]+", text))
    if has_kem_theo and not has_phu_luc:
        return WARN, "VB nhắc 'kèm theo' nhưng không tìm thấy PHỤ LỤC trong file"
    if has_phu_luc:
        return OK, "Có Phụ lục"
    return OK, "Không có phụ lục (không bắt buộc)"


def main():
    if len(sys.argv) < 2:
        print("Usage: python validate_thethuc.py <file.docx>", file=sys.stderr)
        return 1

    path = Path(sys.argv[1]).resolve()
    if not path.is_file():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 1

    doc = Document(str(path))
    text = collect_all_text(doc)

    checks = [
        ("1. Quốc hiệu + Tiêu ngữ", check_quoc_hieu(text)),
        ("2. Tên cơ quan ban hành", check_co_quan(text)),
        ("3. Số/ký hiệu", check_so_van_ban(text)),
        ("4. Tên loại + Trích yếu", check_ten_loai(text)),
        ("5. Nội dung", check_noi_dung(text)),
        ("6. Người ký (chức vụ + tên)", check_nguoi_ky(text)),
        ("7. Dấu/chữ ký số", check_dau()),
        ("8. Nơi nhận + Lưu", check_noi_nhan(text)),
        ("9. Phụ lục", check_phu_luc(text)),
    ]

    print(f"=== Validate thể thức: {path.name} ===\n")
    ok = warn = fail = 0
    for label, (status, detail) in checks:
        print(f"  {status} {label}: {detail}")
        if status == OK: ok += 1
        elif status == WARN: warn += 1
        else: fail += 1

    print(f"\nTổng: {OK}{ok}  {WARN}{warn}  {FAIL}{fail}  /  9 mục")
    if fail > 0:
        return 2
    if warn > 0:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
